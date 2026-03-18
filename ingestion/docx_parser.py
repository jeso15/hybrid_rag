"""Word document (.docx) parser."""

from docx import Document


def parse_docx(file_path: str) -> list[str]:
    """Extract text from a .docx file, one string per paragraph."""
    doc = Document(file_path)
    pages = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            pages.append(text)

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                pages.append(row_text)

    return pages
